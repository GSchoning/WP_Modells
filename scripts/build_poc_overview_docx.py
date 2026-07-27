"""Rebuild docs/gabora-poc-overview-draft.docx from its markdown mirror.

The .md file is the editable source (diffable in git); this script applies
the house report style: Calibri 11 body, navy 1F3864 headings (16/13 pt
bold), A4, bordered disclaimer box, gridded tables. Run after editing the
markdown:

    python scripts/build_poc_overview_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)
SRC = Path(__file__).resolve().parents[1] / "docs" / "gabora-poc-overview-draft.md"
OUT = SRC.with_suffix(".docx")


def parse(md: str):
    """Yield (kind, payload) blocks: title/meta/box/note/h1/h2/p/b/cap/table."""
    lines = md.splitlines()
    i = 0
    table: list[list[str]] = []
    while i < len(lines):
        ln = lines[i].rstrip()
        if ln.startswith("|"):
            cells = [c.strip() for c in ln.strip("|").split("|")]
            if not all(set(c) <= {"-"} for c in cells):      # skip separator row
                table.append(cells)
            i += 1
            continue
        if table:
            yield "table", table
            table = []
        if not ln:
            pass
        elif ln.startswith("# "):
            yield "title", ln[2:]
        elif ln.startswith("## "):
            yield "h1", ln[3:]
        elif ln.startswith("### "):
            yield "h2", ln[4:]
        elif ln.startswith("> "):
            yield "box", ln[2:]
        elif ln.startswith("- "):
            yield "b", ln[2:]
        elif ln.startswith("**Table"):
            yield "cap", ln.strip("*")
        elif ln.startswith("**") and "·" in ln:
            yield "meta", ln.replace("**", "")
        elif ln.startswith("*Citation:*") or ln.startswith("*Contributors:*"):
            yield "note", ln.replace("*", "")
        elif ln == "---":
            yield "pagebreak", ""
        else:
            yield "p", ln
        i += 1
    if table:
        yield "table", table


doc = Document()
for section in doc.sections:
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin = section.bottom_margin = Cm(2.5)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = NAVY
    if st.element.rPr is not None and st.element.rPr.rFonts is not None:
        for attr in ("w:asciiTheme", "w:ascii", "w:hAnsiTheme", "w:hAnsi"):
            st.element.rPr.rFonts.set(qn(attr), "Calibri")


def para(text, style=None, align=None, bold=False, size=None, color=None, after=6):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    return p


for kind, payload in parse(SRC.read_text()):
    if kind == "title":
        para(payload, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18,
             color=NAVY, after=10)
    elif kind == "meta":
        para(payload, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    elif kind == "box":
        box = doc.add_table(rows=1, cols=1)
        box.style = "Table Grid"
        box.rows[0].cells[0].text = payload
        doc.add_paragraph()
    elif kind == "note":
        para(payload, size=9)
    elif kind == "pagebreak":
        doc.add_page_break()
    elif kind == "h1":
        doc.add_heading(payload, level=1)
    elif kind == "h2":
        doc.add_heading(payload, level=2)
    elif kind == "p":
        para(payload)
    elif kind == "b":
        para(payload, style="List Bullet")
    elif kind == "cap":
        para(payload, bold=True, size=10, after=3)
    elif kind == "table":
        t = doc.add_table(rows=len(payload), cols=len(payload[0]))
        t.style = "Table Grid"
        for ri, row in enumerate(payload):
            for ci, val in enumerate(row):
                cell = t.rows[ri].cells[ci]
                cell.text = val
                if ri == 0 and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

doc.save(OUT)
print("saved", OUT)
